import copy
import pandas as pd
import docx
from pathlib import Path
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn
import time

planilha_df = pd.read_excel('base.xlsx')



def tratativa_data(data):
    data_tratada = str(data).split(' ')[0]
    data_tratada = data_tratada.split('-')
    return f'{data_tratada[2]}/{data_tratada[1]}/{data_tratada[0]}'


def tratativa_dataExtenso(data):
    def conversao(data):
        return data[1].replace(data[1], tratamento[str(data[1])])
    
    
    tratamento = {
            '01':'Janeiro',
            '02':'Fevereiro',
            '03':'Março',
            '04':'Abril',
            '05':'Maio',
            '06':'Junho',
            '07':'Julho',
            '08':'Agosto',
            '09':'Setembro',
            '10':'Outubro',
            '11':'Novembro',
            '12':'Dezembro',
            }

    
    data_tratada = str(data).split(' ')[0]
    data_tratada = data_tratada.split('-')
    return f'{data_tratada[2]} de {conversao(data_tratada)} de {data_tratada[0]}'
    


def CriarDocumento(doc_base, nome: str, dados:dict):
    doc_base = copy.deepcopy(doc_base)   
    
    for table in doc_base.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in dados.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))
                        cell.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.name = 'Times New Roman'
                                



    for paragrafo in doc_base.paragraphs:
            if '==DEPOISDEZDIA==' in paragrafo.text:
                paragrafo.text = paragrafo.text.replace('==DEPOISDEZDIA==', dados["==DEPOISDEZDIA=="])

                
    data_salvar = dados["==DESLIGDATA=="].replace('/','.')
    doc_base.save(Path(f'docs feitos/ {data_salvar} - {nome} - TERMINO DE CONTRATO.docx'))
    

    
def run():
    doc_base = docx.Document(r"base.docx")  # Carregar o documento base uma vez
    

    for index, row in planilha_df.iterrows():
        nome = row['NOME']
        admissao = row['ADM']
        lotacao = row['LOTACAO']
        cargo = row['CARGO']
        ctps = row['CTPS']
        dataDesligamento = row['Data Desligamento']
        dataExtenso = row['DATA EXTENSO']
        dataDezDiasDepois = row['DEPOIS DEZ DIAS']
        
        relacao: dict = {
            '==NOME==' : nome,
            '==ADM==' : tratativa_data(admissao),
            '==CARGO==' : cargo,
            '==CTPS==' : ctps,
            '==DEPOISDEZDIA==' : tratativa_data(dataDezDiasDepois),
            '==DESLIGDATA==' : tratativa_data(dataDesligamento),
            '==LOTACAO==' : lotacao.title(),
            '==DESLIGDATAEXTENSO==' : (dataExtenso)
        }
        
        print(relacao)
        CriarDocumento(doc_base, nome, relacao)  # Usar a cópia independente do documento base

run()